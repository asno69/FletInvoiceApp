import os
from datetime import datetime
import flet as ft
from docx import Document
from docx2pdf import convert


def main(page: ft.Page):
    page.title = "Rechnungserstellung"
    page.padding = 10
    page.spacing = 10

    invoice_number_input = ft.TextField(label='Rechnungsnummer')
    date_input = ft.TextField(label='Rechnungsdatum (TT.MM.JJJJ)')
    service_input = ft.TextField(label='Leistungszeitraum')
    salary1_input = ft.TextField(label='Commessa CM0184-003 Assistenza di Commessa Costr 720')
    salary2_input = ft.TextField(label='Commessa CM0189-003 Assistenza di Commessa Costr 718')
    salary3_input = ft.TextField(label='Commessa CM0231-003 Assistenza di Commessa Costr 721')

    def create_invoice(e):
        if any(field.value == '' for field in [salary1_input, salary2_input, salary3_input, invoice_number_input, date_input, service_input]):
            page.add(ft.Text("Alle Werte müssen gesetzt sein"))
            return

        try:
            salary1 = float(salary1_input.value)
            salary2 = float(salary2_input.value)
            salary3 = float(salary3_input.value)
        except ValueError:
            page.add(ft.Text("Bitte geben Sie gültige Zahlen für die Gehälter ein"))
            return

        try:
            date = datetime.strptime(date_input.value, '%d.%m.%Y').strftime('%d.%m.%Y')
        except ValueError:
            page.add(ft.Text('Ungültiges Datum. Bitte im Format TT.MM.JJJJ eingeben.'))
            return

        try:
            template_path = os.path.join(os.path.dirname(__file__), 'Vorlage.docx')
            output_docx_path = os.path.join(os.path.dirname(__file__), f'Rechnung_Nr{invoice_number_input.value}.docx')
            output_pdf_path = output_docx_path.replace('.docx', '.pdf')

            # Load template document
            doc = Document(template_path)

            # Replace placeholders with actual values
            placeholders = {
                '{DATE}': date,
                '{INVOICE_NUMBER}': invoice_number_input.value,
                '{SERVICE}': service_input.value,
                '{SALARY1}': f'{salary1:.2f}',
                '{SALARY2}': f'{salary2:.2f}',
                '{SALARY3}': f'{salary3:.2f}',
                '{BRUTTO}': f'{salary1 + salary2 + salary3:.2f}',
                '{STEUER}': f'{(salary1 + salary2 + salary3) * 0.19:.2f}',
                '{NETTO}': f'{(salary1 + salary2 + salary3) * 0.81:.2f}'
            }

            for paragraph in doc.paragraphs:
                for placeholder, value in placeholders.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, value)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in placeholders.items():
                                if placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, value)

            doc.save(output_docx_path)

            # Convert DOCX to PDF using docx2pdf
            convert(output_docx_path)

            page.add(ft.Text(f'Rechnung wurde erfolgreich erstellt und unter {output_docx_path} und {output_pdf_path} gespeichert!'))

        except Exception as ex:
            page.add(ft.Text(f'Fehler beim Erstellen der Rechnung: {ex}'))

    submit_button = ft.ElevatedButton(text='Rechnung erstellen', on_click=create_invoice)
    page.add(invoice_number_input, date_input, service_input, salary1_input, salary2_input, salary3_input, submit_button)


# Start der Anwendung
ft.app(target=main)
